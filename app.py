import sqlite3
import io
import os
from flask import Flask, render_template, redirect, request, flash, send_file
from werkzeug.exceptions import abort
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

app = Flask(__name__)
app.config['SECRET_KEY'] = os.urandom(24)
UPLOAD_FOLDER = 'contracts'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def get_db_connection():
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    return conn

def _generate_contract_docx(contract):
    try:
        doc = Document()

        title = doc.add_heading('CYBER ARENA COMPUTER CLUB', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('ЧЕК ОПЛАТЫ', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()

        doc.add_heading('ИНФОРМАЦИЯ О СЕАНСЕ:', level=2)
        
        p = doc.add_paragraph()
        p.add_run('Номер сеанса: ').bold = True
        p.add_run(f"#{contract.get('numbers', 'N/A')}")
        
        p = doc.add_paragraph()
        p.add_run('Дата и время: ').bold = True
        p.add_run(f"{contract.get('dates', 'N/A')}")
        
        p = doc.add_paragraph()
        p.add_run('Тариф в час: ').bold = True
        p.add_run(f"{contract.get('price', 'N/A')} руб/час")
        
        p = doc.add_paragraph()
        p.add_run('Скидка: ').bold = True
        p.add_run(f"{contract.get('discount', 0)}%")
        
        p = doc.add_paragraph()
        p.add_run('Итоговая стоимость: ').bold = True
        p.add_run(f"{contract.get('finish_price', 'N/A')} руб")
        
        doc.add_paragraph()

        doc.add_heading('ИНФОРМАЦИЯ О КЛИЕНТЕ:', level=2)
        
        p = doc.add_paragraph()
        p.add_run('ФИО: ').bold = True
        p.add_run(f"{contract.get('client_name', 'N/A')}")
        
        p = doc.add_paragraph()
        p.add_run('Телефон: ').bold = True
        p.add_run(f"{contract.get('client_phone', 'N/A')}")
        
        p = doc.add_paragraph()
        p.add_run('Email: ').bold = True
        p.add_run(f"{contract.get('client_email', 'Не указан')}")
        
        doc.add_paragraph()

        doc.add_heading('ИНФОРМАЦИЯ О КОМПЬЮТЕРЕ:', level=2)
        
        p = doc.add_paragraph()
        p.add_run('Местоположение: ').bold = True
        p.add_run(f"{contract.get('device_address', 'N/A')}")
        
        p = doc.add_paragraph()
        p.add_run('Филиал: ').bold = True
        p.add_run(f"{contract.get('branch_name', 'N/A')}")
        
        p = doc.add_paragraph()
        p.add_run('Сотрудник: ').bold = True
        p.add_run(f"{contract.get('employee_name', 'N/A')}")
        
        p = doc.add_paragraph()
        p.add_run('Статус устройства: ').bold = True
        p.add_run(f"{contract.get('device_status', 'Норма')}")
        
        doc.add_paragraph()

        footer = doc.add_paragraph()
        footer_run = footer.add_run(f"Чек сгенерирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        footer_run.italic = True
        
        thanks = doc.add_paragraph('Спасибо за посещение!')
        thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        print(f"Ошибка в _generate_contract_docx: {e}")
        import traceback
        traceback.print_exc()
        return None

def _generate_client_report_docx(client):
    try:
        doc = Document()

        title = doc.add_heading('CYBER ARENA COMPUTER CLUB', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('ОТЧЕТ ПО КЛИЕНТУ', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()

        doc.add_heading('ОСНОВНАЯ ИНФОРМАЦИЯ:', level=2)
        
        p = doc.add_paragraph()
        p.add_run('Клиент: ').bold = True
        p.add_run(f"{client.get('name', 'N/A')}")
        
        p = doc.add_paragraph()
        p.add_run('Email: ').bold = True
        p.add_run(f"{client.get('email', 'Не указан')}")
        
        p = doc.add_paragraph()
        p.add_run('Телефон: ').bold = True
        p.add_run(f"{client.get('phone_number', 'N/A')}")
        
        doc.add_paragraph()
        
        doc.add_heading('ИСТОРИЯ СЕАНСОВ:', level=2)
        
        if client.get('sessions'):
            for i, session in enumerate(client['sessions'], 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(f"Дата: {session.get('dates', '')}, ")
                p.add_run(f"Компьютер: {session.get('device_address', '')}, ")
                p.add_run(f"Стоимость: {session.get('finish_price', 0)} руб")
        else:
            p = doc.add_paragraph()
            p_run = p.add_run('История сеансов: нет данных')
            p_run.italic = True
        
        doc.add_paragraph()

        footer = doc.add_paragraph()
        footer_run = footer.add_run(f"Отчет сгенерирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        footer_run.italic = True

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        print(f"Ошибка в _generate_client_report_docx: {e}")
        import traceback
        traceback.print_exc()
        return None

@app.route('/generate_client_report/<int:client_id>')
def generate_client_report(client_id):
    try:
        client = get_client(client_id)
        
        if client is None:
            flash('Клиент не найден')
            return redirect('/clients')

        docx_buffer = _generate_client_report_docx(client)
        
        if docx_buffer is None:
            flash('Ошибка при генерации отчета')
            return redirect('/clients')

        filename = f"отчет_клиент_{client.get('name', 'unknown')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        flash(f'Ошибка при генерации отчета клиента: {str(e)}')
        return redirect('/clients')

@app.route('/')
def index():
    return redirect("/contracts")

@app.route('/contracts')
def contracts():
    conn = get_db_connection()
    contracts = conn.execute("""SELECT contracts.*, 
                                       client.name as client_name, 
                                       client.phone_number as client_phone,
                                       device.address as device_address,
                                       branch.branch_name as branch_name,
                                       employee.name as employee_name
                                FROM contracts 
                                LEFT JOIN client ON contracts.client_id = client.id
                                LEFT JOIN device ON contracts.device_id = device.id
                                LEFT JOIN employee ON contracts.employee_id = employee.id
                                LEFT JOIN branch ON device.branch_id = branch.id
                                ORDER BY contracts.dates DESC
                             """).fetchall()
    conn.close()
    return render_template('contracts.html', contracts=contracts)

def get_contract(contract_id):
    conn = get_db_connection()
    contract = conn.execute("""SELECT contracts.*, 
                                      client.name as client_name, 
                                      client.email as client_email,
                                      client.phone_number as client_phone,
                                      device.address as device_address,
                                      branch.branch_name as branch_name,
                                      employee.name as employee_name
                               FROM contracts 
                               LEFT JOIN client ON contracts.client_id = client.id
                               LEFT JOIN device ON contracts.device_id = device.id
                               LEFT JOIN employee ON contracts.employee_id = employee.id
                               LEFT JOIN branch ON device.branch_id = branch.id
                               WHERE contracts.id = ?
                            """, (contract_id,)).fetchone()
    conn.close()
    if contract is None:
        abort(404)
    return contract

@app.route('/contract/<int:contract_id>')
def contract(contract_id):
    contract = get_contract(contract_id)
    return render_template('contract.html', contract=contract)

@app.route('/new_contract', methods=('GET', 'POST'))
def new_contract():
    if request.method == 'POST':
        try:
            number = request.form['number']
            date = request.form['date']
            price = int(request.form['price'])
            discount = int(request.form.get('discount', 0))
            device_status = request.form.get('device_status', 'Норма')
            client_id = int(request.form.get('client_id'))
            device_id = int(request.form.get('device_id'))
            employee_id = int(request.form.get('employee_id'))
            finish_price = price * (100 - discount) // 100
        except (ValueError, KeyError) as e:
            flash(f'Некорректные значения: {str(e)}')
            return render_template('new_contract.html', 
                                 clients=get_clients(), 
                                 devices=get_devices(), 
                                 employees=get_employees())
        
        if not (number and date and price and client_id and device_id and employee_id):
            flash('Не все обязательные поля заполнены')
        else:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("""INSERT INTO contracts 
                             (numbers, dates, price, discount, device_status, device_id, client_id, employee_id, finish_price)  
                             VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (number, date, price, discount, device_status, device_id, client_id, employee_id, finish_price))
            conn.commit()
            new_contract_id = cursor.lastrowid
            conn.close()
            return redirect(f'/contract/{new_contract_id}')
    
    return render_template('new_contract.html', 
                         clients=get_clients(), 
                         devices=get_devices(), 
                         employees=get_employees())

@app.route('/generate_contract')
def generate_contract():
    try:
        contract_id = request.args.get('contract_id')
        
        if not contract_id:
            flash('Не указан ID контракта')
            return redirect('/contracts')
        
        conn = get_db_connection()
        contract = conn.execute("""
            SELECT contracts.*, 
                   client.name as client_name, 
                   client.email as client_email,
                   client.phone_number as client_phone,
                   device.address as device_address,
                   branch.branch_name as branch_name,
                   employee.name as employee_name
            FROM contracts 
            LEFT JOIN client ON contracts.client_id = client.id
            LEFT JOIN device ON contracts.device_id = device.id
            LEFT JOIN employee ON contracts.employee_id = employee.id
            LEFT JOIN branch ON device.branch_id = branch.id
            WHERE contracts.id = ?
        """, (contract_id,)).fetchone()
        conn.close()
        
        if contract is None:
            flash('Контракт не найден')
            return redirect('/contracts')

        contract_dict = dict(contract)

        docx_buffer = _generate_contract_docx(contract_dict)
        
        if docx_buffer is None:
            flash('Ошибка при генерации файла')
            return redirect('/contracts')

        filename = f"чек_сеанс_{contract_dict.get('numbers', 'unknown')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        flash(f'Ошибка при генерации чека: {str(e)}')
        return redirect('/contracts')


###############
###############
############### Компьютеры (устройства)
###############
###############

@app.route('/devices')
def devices():
    conn = get_db_connection()
    devices = conn.execute("""SELECT device.*, 
                                     branch.branch_name,
                                     client.name as current_client_name
                              FROM device 
                              LEFT JOIN branch ON device.branch_id = branch.id
                              LEFT JOIN client ON device.client_id = client.id
                           """).fetchall()
    conn.close()
    return render_template('devices.html', devices=devices)


def get_device(device_id):
    conn = get_db_connection()
    device = conn.execute("""SELECT device.*, 
                                    branch.branch_name,
                                    client.name as current_client_name
                             FROM device 
                             LEFT JOIN branch ON device.branch_id = branch.id
                             LEFT JOIN client ON device.client_id = client.id
                             WHERE device.id = ?
                          """, (device_id,)).fetchone()
    
    sessions = conn.execute("""SELECT contracts.*, client.name as client_name
                              FROM contracts 
                              LEFT JOIN client ON contracts.client_id = client.id
                              WHERE contracts.device_id = ?
                              ORDER BY contracts.dates DESC
                              LIMIT 10
                           """, (device_id,)).fetchall()
    conn.close()
    
    if device is None:
        abort(404)

    device_dict = dict(device)
    device_dict['session_history'] = [dict(session) for session in sessions]
    return device_dict


@app.route('/device/<int:device_id>')
def device(device_id):
    device = get_device(device_id)
    return render_template('device.html', device=device)


@app.route('/new_device', methods=('GET', 'POST'))
def new_device():
    if request.method == 'POST':
        address = request.form['address']
        branch_id = request.form.get('branch_id')
        
        if not (address and branch_id):
            flash('Не все поля заполнены')
        else:
            conn = get_db_connection()
            conn.execute("INSERT INTO device (address, branch_id) VALUES (?, ?)",
                        (address, branch_id))
            conn.commit()
            conn.close()
            return redirect('/devices')
    
    conn = get_db_connection()
    branches = conn.execute("SELECT * FROM branch").fetchall()
    conn.close()
    return render_template('new_device.html', branches=branches)


###############
###############
############### Клиенты
###############
###############

def get_clients():
    conn = get_db_connection()
    clients = conn.execute("SELECT * FROM client").fetchall()
    conn.close()
    return clients

def get_client(client_id):
    conn = get_db_connection()
    client = conn.execute("SELECT * FROM client WHERE id = ?", (client_id,)).fetchone()

    sessions = conn.execute("""SELECT contracts.*, device.address as device_address, branch.branch_name
                             FROM contracts 
                             LEFT JOIN device ON contracts.device_id = device.id
                             LEFT JOIN branch ON device.branch_id = branch.id
                             WHERE contracts.client_id = ?
                             ORDER BY contracts.dates DESC
                             LIMIT 10
                          """, (client_id,)).fetchall()
    conn.close()
    
    if client is None:
        abort(404)
    
    client_dict = dict(client)
    client_dict['sessions'] = [dict(session) for session in sessions]
    return client_dict

@app.route('/clients')
def clients():
    conn = get_db_connection()
    clients = conn.execute("SELECT * FROM client").fetchall()
    conn.close()
    return render_template('clients.html', clients=clients)

@app.route('/client/<int:client_id>')
def client(client_id):
    client = get_client(client_id)
    return render_template('client.html', client=client)

@app.route('/new_client', methods=('GET', 'POST'))
def new_client():
    if request.method == 'POST':
        name = request.form['name']
        email = request.form['email']
        phone_number = request.form['phone_number']
        
        if not (name and phone_number):
            flash('Не все обязательные поля заполнены')
        else:
            conn = get_db_connection()
            conn.execute("INSERT INTO client (name, email, phone_number) VALUES (?, ?, ?)",
                        (name, email, phone_number))
            conn.commit()
            conn.close()
            return redirect('/clients')
    
    return render_template('new_client.html')


###############
###############
############### Сотрудники
###############
###############

def get_employees():
    conn = get_db_connection()
    employees = conn.execute("SELECT * FROM employee").fetchall()
    conn.close()
    return employees

def get_employee(employee_id):
    conn = get_db_connection()
    employee = conn.execute("""SELECT employee.*, 
                                     branch.branch_name,
                                     chief.name as chief_name
                              FROM employee 
                              LEFT JOIN branch ON employee.branch_id = branch.id
                              LEFT JOIN employee as chief ON employee.chief_id = chief.id
                              WHERE employee.id = ?
                           """, (employee_id,)).fetchone()
    conn.close()
    if employee is None:
        abort(404)
    return employee

@app.route('/employees')
def employees():
    conn = get_db_connection()
    employees = conn.execute("""SELECT employee.*, branch.branch_name
                              FROM employee 
                              LEFT JOIN branch ON employee.branch_id = branch.id
                           """).fetchall()
    conn.close()
    return render_template('employees.html', employees=employees)

@app.route('/employee/<int:employee_id>')
def employee(employee_id):
    employee = get_employee(employee_id)
    return render_template('employee.html', employee=employee)

@app.route('/new_employee', methods=('GET', 'POST'))
def new_employee():
    if request.method == 'POST':
        name = request.form['name']
        email = request.form['email']
        phone_number = request.form['phone_number']
        position = request.form['position']
        chief_id = request.form.get('chief_id') or None
        branch_id = request.form.get('branch_id')
        
        if not (name and phone_number and position and branch_id):
            flash('Не все обязательные поля заполнены')
        else:
            conn = get_db_connection()
            conn.execute("""INSERT INTO employee 
                         (name, email, phone_number, position, chief_id, branch_id) 
                         VALUES (?, ?, ?, ?, ?, ?)""",
                        (name, email, phone_number, position, chief_id, branch_id))
            conn.commit()
            conn.close()
            return redirect('/employees')
    
    conn = get_db_connection()
    chiefs = conn.execute("SELECT id, name FROM employee").fetchall()
    branches = conn.execute("SELECT id, branch_name FROM branch").fetchall()
    conn.close()
    return render_template('new_employee.html', chiefs=chiefs, branches=branches)


###############
###############
############### Филиалы
###############
###############

@app.route('/branches')
def branches():
    conn = get_db_connection()
    branches = conn.execute("SELECT * FROM branch").fetchall()
    conn.close()
    return render_template('branches.html', branches=branches)

@app.route('/branch/<int:branch_id>')
def branch(branch_id):
    conn = get_db_connection()
    branch = conn.execute("SELECT * FROM branch WHERE id = ?", (branch_id,)).fetchone()

    devices = conn.execute("SELECT * FROM device WHERE branch_id = ?", (branch_id,)).fetchall()

    employees = conn.execute("SELECT * FROM employee WHERE branch_id = ?", (branch_id,)).fetchall()
    conn.close()
    
    if branch is None:
        abort(404)
    
    branch_dict = dict(branch)
    branch_dict['devices'] = [dict(device) for device in devices]
    branch_dict['employees'] = [dict(employee) for employee in employees]
    return render_template('branch.html', branch=branch_dict)

@app.route('/new_branch', methods=('GET', 'POST'))
def new_branch():
    if request.method == 'POST':
        address = request.form['address']
        branch_name = request.form['branch_name']
        
        if not (address and branch_name):
            flash('Не все поля заполнены')
        else:
            conn = get_db_connection()
            conn.execute("INSERT INTO branch (address, branch_name) VALUES (?, ?)",
                        (address, branch_name))
            conn.commit()
            conn.close()
            return redirect('/branches')
    
    return render_template('new_branch.html')


###############
###############
############### Отчеты
###############
###############

@app.route('/reports')
def reports():
    conn = get_db_connection()
    reports = conn.execute("""SELECT reports.*, employee.name as employee_name
                            FROM reports 
                            LEFT JOIN employee ON reports.employee_id = employee.id
                            ORDER BY reports.date DESC
                         """).fetchall()
    conn.close()
    return render_template('reports.html', reports=reports)

@app.route('/report/<int:report_id>')
def report(report_id):
    conn = get_db_connection()
    report = conn.execute("""SELECT reports.*, employee.name as employee_name
                           FROM reports 
                           LEFT JOIN employee ON reports.employee_id = employee.id
                           WHERE reports.id = ?
                        """, (report_id,)).fetchone()
    conn.close()
    if report is None:
        abort(404)
    return render_template('report.html', report=report)

@app.route('/new_report', methods=('GET', 'POST'))
def new_report():
    if request.method == 'POST':
        number = request.form['number']
        date = request.form['date']
        report_type = request.form['report_type']
        description = request.form['description']
        employee_id = request.form.get('employee_id')
        
        if not (number and date and report_type and employee_id):
            flash('Не все обязательные поля заполнены')
        else:
            conn = get_db_connection()
            conn.execute("""INSERT INTO reports 
                         (number, date, report_type, description, employee_id) 
                         VALUES (?, ?, ?, ?, ?)""",
                        (number, date, report_type, description, employee_id))
            conn.commit()
            conn.close()
            return redirect('/reports')
    
    conn = get_db_connection()
    employees = conn.execute("SELECT id, name FROM employee").fetchall()
    conn.close()
    return render_template('new_report.html', employees=employees)


###############
###############
############### Вспомогательные функции
###############
###############

def get_devices():
    conn = get_db_connection()
    devices = conn.execute("SELECT * FROM device").fetchall()
    conn.close()
    return devices

#
# 404 
#

@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404


if __name__ == '__main__':
    app.run(debug=True)